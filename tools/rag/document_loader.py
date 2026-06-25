# tools/rag/document_loader.py

"""
文档加载器。

支持 PDF, DOCX, TXT, Markdown 文件，以及网页爬取。
返回纯文本内容。

依赖（可选）:
    pip install python-docx PyPDF2 beautifulsoup4
"""

import os
import logging
from typing import Optional
from dataclasses import dataclass, field

logger = logging.getLogger(__name__)


@dataclass
class LoadedDocument:
    """加载后的文档"""
    content: str
    source: str                # 文件名或 URL
    file_type: str             # pdf / docx / txt / md / html
    metadata: dict = field(default_factory=dict)


class DocumentLoader:
    """文档加载器"""

    SUPPORTED_EXTENSIONS = {".txt", ".md", ".pdf", ".docx", ".html", ".htm"}

    async def load_file(self, path: str) -> LoadedDocument:
        """从本地文件加载"""
        if not os.path.exists(path):
            raise FileNotFoundError(f"File not found: {path}")
        ext = os.path.splitext(path)[1].lower()
        if ext not in self.SUPPORTED_EXTENSIONS:
            raise ValueError(f"Unsupported file type: {ext}")

        content = ""
        if ext in (".txt", ".md"):
            content = self._load_text(path)
        elif ext == ".pdf":
            content = self._load_pdf(path)
        elif ext == ".docx":
            content = self._load_docx(path)
        elif ext in (".html", ".htm"):
            content = self._load_html(path)

        return LoadedDocument(
            content=content,
            source=path,
            file_type=ext.lstrip("."),
            metadata={"filename": os.path.basename(path), "size": os.path.getsize(path)},
        )

    async def load_url(self, url: str) -> LoadedDocument:
        """从网页爬取内容"""
        try:
            import httpx
            from bs4 import BeautifulSoup
        except ImportError:
            raise ImportError("pip install httpx beautifulsoup4")

        async with httpx.AsyncClient(timeout=30.0, follow_redirects=True) as client:
            resp = await client.get(url)
            resp.raise_for_status()
            soup = BeautifulSoup(resp.text, "html.parser")
            # 移除 script/style
            for tag in soup(["script", "style", "nav", "footer", "header"]):
                tag.decompose()
            content = soup.get_text(separator="\n", strip=True)

        return LoadedDocument(
            content=content,
            source=url,
            file_type="html",
            metadata={"url": url, "title": soup.title.string if soup.title else ""},
        )

    async def load_stream(self, data: bytes, filename: str) -> LoadedDocument:
        """从字节流加载（用于上传接口）"""
        ext = os.path.splitext(filename)[1].lower()
        if ext == ".txt" or ext == ".md":
            content = data.decode("utf-8")
        elif ext == ".pdf":
            # 写入临时文件再加载
            import tempfile
            with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as tmp:
                tmp.write(data)
                tmp_path = tmp.name
            try:
                result = await self.load_file(tmp_path)
                result.source = filename
                return result
            finally:
                os.unlink(tmp_path)
        else:
            content = data.decode("utf-8", errors="replace")

        return LoadedDocument(content=content, source=filename, file_type=ext.lstrip("."))

    def _load_text(self, path: str) -> str:
        with open(path, "r", encoding="utf-8", errors="replace") as f:
            return f.read()

    def _load_html(self, path: str) -> str:
        try:
            from bs4 import BeautifulSoup
            with open(path, "r", encoding="utf-8") as f:
                soup = BeautifulSoup(f.read(), "html.parser")
            return soup.get_text(separator="\n", strip=True)
        except ImportError:
            return self._load_text(path)

    def _load_pdf(self, path: str) -> str:
        try:
            import PyPDF2
            reader = PyPDF2.PdfReader(path)
            pages = []
            for i, page in enumerate(reader.pages):
                text = page.extract_text()
                if text:
                    pages.append(text)
            return "\n\n".join(pages)
        except ImportError:
            raise ImportError("pip install PyPDF2 for PDF support")

    def _load_docx(self, path: str) -> str:
        try:
            from docx import Document
            doc = Document(path)
            return "\n".join(p.text for p in doc.paragraphs if p.text.strip())
        except ImportError:
            raise ImportError("pip install python-docx for DOCX support")
